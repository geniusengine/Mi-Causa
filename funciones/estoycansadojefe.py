
import sys
import os
from PyQt6.QtWidgets import QApplication, QMainWindow, QTextEdit, QPushButton, QVBoxLayout, QWidget, QFileDialog
import mysql.connector
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import smtplib


class WordDocumentApp(QMainWindow):
    def __init__(self):
        super().__init__()

        self.initUI()
        self.file_name = None

    def initUI(self):
        self.setWindowTitle("Editor de Documentos de Word")
        self.setGeometry(100, 100, 800, 600)

        self.text_edit = QTextEdit(self)
        self.text_edit.setAcceptRichText(True)

        self.load_button = QPushButton("Cargar Documento de Word", self)
        self.load_button.clicked.connect(self.load_word_doc)

        self.save_button = QPushButton("Guardar Cambios", self)
        self.save_button.clicked.connect(self.save_document)

        self.print_button = QPushButton("Imprimir Documento", self)
        self.print_button.clicked.connect(self.print_document)

        self.email_button = QPushButton("Enviar por Correo", self)
        self.email_button.clicked.connect(self.send_email)

        self.stamp_button = QPushButton("Agregar Estampado", self)
        self.stamp_button.clicked.connect(self.add_stamp)

        self.fill_button = QPushButton("Llenar con Datos", self)
        self.fill_button.clicked.connect(self.fill_with_data)

        layout = QVBoxLayout()
        layout.addWidget(self.load_button)
        layout.addWidget(self.save_button)
        layout.addWidget(self.text_edit)
        layout.addWidget(self.print_button)
        layout.addWidget(self.email_button)
        layout.addWidget(self.stamp_button)
        layout.addWidget(self.fill_button)

        container = QWidget()
        container.setLayout(layout)

        self.setCentralWidget(container)

        self.db = mysql.connector.connect(
            host="localhost",
            user="root",
            password="",
            database="mi_causa"
        )

    def load_word_doc(self):
        self.file_name, _ = QFileDialog.getOpenFileName(self, "Abrir Documento de Word", "", "Archivos de Word (*.docx);;Todos los archivos (*)")
        if self.file_name:
            doc = Document(self.file_name)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text
            self.text_edit.setPlainText(text)

    def save_document(self):
        if self.file_name is not None:
            text = self.text_edit.toPlainText()
            doc = Document()
            for line in text.split('\n'):
                doc.add_paragraph(line)
            doc.save(self.file_name)

    def print_document(self):
        if self.file_name is not None:
            doc = SimpleDocTemplate("print_output.pdf", pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            with open(self.file_name, 'r', encoding='utf-8') as file:
                for line in file:
                    story.append(Paragraph(line, styles["Normal"]))
            doc.build(story)
            os.system("lpr -P YOUR_PRINTER_NAME print_output.pdf")  # Reemplaza YOUR_PRINTER_NAME

    def send_email(self):
        if self.file_name is not None:
            from_email = "tu_correo@gmail.com"  # Tu dirección de correo
            password = "tu_contraseña"  # Tu contraseña
            recipient_email = "destinatario@gmail.com"  # Correo del destinatario

            subject = "Documento adjunto"
            msg = MIMEMultipart()
            msg["From"] = from_email
            msg["To"] = recipient_email
            msg["Subject"] = subject

            with open(self.file_name, "rb") as file:
                part = MIMEApplication(file.read(), Name=os.path.basename(self.file_name))
            part["Content-Disposition"] = f'attachment; filename="{os.path.basename(self.file_name)}"'
            msg.attach(part)

            server = smtplib.SMTP("smtp.gmail.com", 587)
            server.starttls()
            server.login(from_email, password)
            server.sendmail(from_email, recipient_email, msg.as_string())
            server.quit()
    

  

    def add_stamp(self):

        doc = Document(self.file_name)

        # Agregar un nuevo párrafo al documento.
        paragraph = doc.add_paragraph()

        # Agregar un nuevo fragmento de texto al párrafo.
        run = paragraph.add_run()

        # Agregar la imagen del estampado al fragmento de texto.
        run.add_picture("recursos\hola.png")

        # Regresar el documento con el estampado agregado.
        return doc



    def fill_with_data(self):
        if self.file_name is not None:
            cursor = self.db.cursor()
            select_query = "SELECT nombre, apellido, edad FROM datoscausas"
            cursor.execute(select_query)
            data = cursor.fetchone()
            cursor.close()

            doc = Document(self.file_name)
            for paragraph in doc.paragraphs:
                if "Nombre:" in paragraph.text:
                    paragraph.clear()
                    paragraph.add_run(f"Nombre: {data[0]}")
                elif "Apellido:" in paragraph.text:
                    paragraph.clear()
                    paragraph.add_run(f"Apellido: {data[1]}")
                elif "Edad:" in paragraph.text:
                    paragraph.clear()
                    paragraph.add_run(f"Edad: {data[2]}")
            doc.save(self.file_name)

def main():
    app = QApplication(sys.argv)
    window = WordDocumentApp()
    window.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()

